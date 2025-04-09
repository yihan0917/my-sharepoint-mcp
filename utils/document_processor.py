"""Document processing utilities for SharePoint MCP server."""

import io
import logging
from typing import Dict, Any, Optional

# Packages to support different file formats
try:
    import pandas as pd
    import docx
    from PyPDF2 import PdfReader
    import openpyxl
    HAS_DOCUMENT_LIBRARIES = True
except ImportError:
    HAS_DOCUMENT_LIBRARIES = False

# Setup logging
logger = logging.getLogger("document_processor")

class DocumentProcessor:
    """Processor for various document types."""
    
    @staticmethod
    def check_dependencies():
        """Check if all required dependencies are installed."""
        if not HAS_DOCUMENT_LIBRARIES:
            logger.warning("Document processing libraries are not installed.")
            logger.warning("Please install with: pip install pandas python-docx PyPDF2 openpyxl")
            return False
        return True
    
    @staticmethod
    def process_document(content: bytes, filename: str) -> Dict[str, Any]:
        """Process document content based on file type.
        
        Args:
            content: Document content as bytes
            filename: Name of the file
        
        Returns:
            Processed document information
        """
        if not DocumentProcessor.check_dependencies():
            return {"error": "Document processing libraries not installed"}
        
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if file_ext == 'csv':
                return DocumentProcessor._process_csv(content)
            elif file_ext in ('xlsx', 'xls'):
                return DocumentProcessor._process_excel(content)
            elif file_ext == 'docx':
                return DocumentProcessor._process_word(content)
            elif file_ext == 'pdf':
                return DocumentProcessor._process_pdf(content)
            elif file_ext in ('txt', 'md', 'html', 'htm'):
                return DocumentProcessor._process_text(content)
            else:
                return {"error": f"Unsupported file type: {file_ext}"}
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    def _process_csv(content: bytes) -> Dict[str, Any]:
        """Process CSV content.
        
        Args:
            content: CSV file content
            
        Returns:
            Processed data and analysis
        """
        df = pd.read_csv(io.BytesIO(content))
        return {
            "type": "csv",
            "rows": len(df),
            "columns": list(df.columns),
            "preview": df.head(5).to_dict(orient='records'),
            "summary": {
                "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
                "missing_values": df.isnull().sum().to_dict(),
                "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
        }
    
    @staticmethod
    def _process_excel(content: bytes) -> Dict[str, Any]:
        """Process Excel content.
        
        Args:
            content: Excel file content
            
        Returns:
            Processed data and analysis
        """
        df_dict = pd.read_excel(io.BytesIO(content), sheet_name=None)
        sheets = {}
        
        for sheet_name, df in df_dict.items():
            sheets[sheet_name] = {
                "rows": len(df),
                "columns": list(df.columns),
                "preview": df.head(5).to_dict(orient='records'),
                "summary": {
                    "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
                    "missing_values": df.isnull().sum().to_dict(),
                    "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
                }
            }
        
        return {
            "type": "excel",
            "sheet_count": len(sheets),
            "sheets": sheets
        }
    
    @staticmethod
    def _process_word(content: bytes) -> Dict[str, Any]:
        """Process Word document content.
        
        Args:
            content: Word document content
            
        Returns:
            Processed text and structure
        """
        doc = docx.Document(io.BytesIO(content))
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        
        for table in doc.tables:
            t_data = []
            for row in table.rows:
                t_data.append([cell.text for cell in row.cells])
            tables.append(t_data)
        
        # Extract document metadata
        core_properties = {}
        try:
            props = doc.core_properties
            core_properties = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "category": props.category or ""
            }
        except Exception as e:
            logger.warning(f"Error getting document properties: {str(e)}")
        
        # Extract sections and headings for document structure
        structure = []
        heading_styles = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Title'}
        for p in doc.paragraphs:
            if p.style.name in heading_styles and p.text.strip():
                level = 0
                if p.style.name == 'Title':
                    level = 0
                else:
                    # Extract heading level number from style name
                    try:
                        level = int(p.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                
                structure.append({
                    "level": level,
                    "text": p.text
                })
        
        return {
            "type": "word",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "content": paragraphs[:20],  # First 20 paragraphs only
            "tables": tables[:5],  # First 5 tables only
            "properties": core_properties,
            "structure": structure
        }
    
    @staticmethod
    def _process_pdf(content: bytes) -> Dict[str, Any]:
        """Process PDF content.
        
        Args:
            content: PDF file content
            
        Returns:
            Extracted text and metadata
        """
        pdf = PdfReader(io.BytesIO(content))
        
        pages = []
        for i in range(min(10, len(pdf.pages))):  # First 10 pages only
            pages.append(pdf.pages[i].extract_text())
        
        # Extract metadata
        metadata = {}
        if pdf.metadata:
            for key, value in pdf.metadata.items():
                if key.startswith('/'):
                    key = key[1:]  # Remove leading slash
                if isinstance(value, (str, int, float, bool)) and key not in ('Trapped'):
                    metadata[key] = str(value)
        
        # Extract form fields if present
        form_fields = []
        if hasattr(pdf, 'get_fields'):
            fields = pdf.get_fields()
            if fields:
                for field_name, field_value in fields.items():
                    form_fields.append({
                        "name": field_name,
                        "value": str(field_value)
                    })
        
        return {
            "type": "pdf",
            "page_count": len(pdf.pages),
            "content": pages,
            "metadata": metadata,
            "form_fields": form_fields
        }
    
    @staticmethod
    def _process_text(content: bytes) -> Dict[str, Any]:
        """Process text content.
        
        Args:
            content: Text file content
            
        Returns:
            Processed text information
        """
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = content.decode('latin-1')
            except Exception:
                return {"error": "Failed to decode text content"}
        
        lines = text.splitlines()
        
        # Calculate some basic statistics
        word_count = len(text.split())
        char_count = len(text)
        avg_line_length = char_count / len(lines) if lines else 0
        
        # Try to detect if it's structured data like markdown or html
        is_markdown = text.count('#') > 0 and text.count('##') > 0
        is_html = text.count('<html') > 0 or text.count('<body') > 0
        
        return {
            "type": "text",
            "line_count": len(lines),
            "word_count": word_count,
            "character_count": char_count,
            "average_line_length": round(avg_line_length, 2),
            "content": lines[:30],  # First 30 lines only
            "format": "html" if is_html else ("markdown" if is_markdown else "plain_text")
        }
